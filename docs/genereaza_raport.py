# -*- coding: utf-8 -*-
"""
Generator de raport pentru proiectul LUMINA.
Dintr-o singură sursă de conținut produce:
  - docs/raport-practica.md   (varianta Markdown, ușor de citit / versionat)
  - docs/Raport-practica-LUMINA.docx  (Word stilizat, cu pagină de titlu + cuprins)
  - docs/Raport-practica-LUMINA.pdf   (PDF, convertit din Word, cu cuprins actualizat)

Rulare:  python docs/genereaza_raport.py
Cerințe: python-docx (Word + PDF necesită Microsoft Word instalat, prin win32com).
"""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------
# Căi
# ----------------------------------------------------------------------
DOCS = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(DOCS)
SHOTS = os.path.join(ROOT, "assets", "images", "screenshots")
MD_OUT = os.path.join(DOCS, "raport-practica.md")
DOCX_OUT = os.path.join(DOCS, "Raport-practica-LUMINA.docx")
PDF_OUT = os.path.join(DOCS, "Raport-practica-LUMINA.pdf")

DEMO_URL = "https://web-practica-one.vercel.app"
REPO_URL = "https://github.com/AlbotCiprian/web_practica"

# Culori din identitatea vizuală a site-ului
MARO = RGBColor(0x3B, 0x2A, 0x20)
CARAMEL = RGBColor(0xC8, 0x89, 0x3F)
CARAMEL_INCHIS = RGBColor(0xA6, 0x6E, 0x2C)
MARO_MEDIU = RGBColor(0x6B, 0x54, 0x44)

# ======================================================================
# 1. CONȚINUTUL RAPORTULUI – sursă unică (listă de blocuri)
#    Tipuri: h1, h2, h3, p, bullets, code, image, pagebreak
# ======================================================================
COVER = {
    "supratitlu": "FACULTATEA DE MATEMATICĂ ȘI INFORMATICĂ",
    "titlu": "RAPORT DE PRACTICĂ",
    "subtitlu": "Aplicație web — LUMINA · Café & Reading Corner",
    "autor": "Nelly Prijilevschi",
    "grupa": "Grupa: IAFR2502",
    "an": "Anul I",
    "loc": "Chișinău, 2026",
    "demo": DEMO_URL,
}

CONTENT = [
    ("h1", "1. Introducere și prezentarea scopului aplicației"),
    ("p", "Prezentul raport descrie proiectul de practică **LUMINA – Café & Reading Corner**, "
          "o aplicație web de prezentare pentru o cafenea boutique care îmbină trei concepte: "
          "cafeaua de specialitate, un colț de lectură și organizarea de evenimente culturale mici "
          "(seri de poezie, cluburi de carte, concerte acustice)."),
    ("p", "Scopul aplicației este dublu. Din perspectiva afacerii, site-ul oferă vizitatorilor o "
          "prezentare atractivă a locației, un meniu interactiv, o listă de evenimente și posibilitatea "
          "de a face o rezervare sau de a trimite un mesaj. Din perspectivă academică, proiectul "
          "demonstrează construirea unui site complet funcțional și responsive folosind **doar** "
          "tehnologiile web fundamentale, fără framework-uri, fără backend și fără bază de date."),
    ("p", "Tema rezolvă o nevoie reală a unui mic business local: o prezență online curată, rapidă și "
          "ușor de folosit atât pe telefon, cât și pe desktop, cu câteva funcționalități interactive "
          "(comandă demonstrativă, favorite, formulare validate) care îmbunătățesc experiența "
          "utilizatorului, fără costuri de infrastructură."),
    ("p", "Aplicația este publicată online și poate fi accesată la adresa: **%s**, "
          "iar codul sursă este disponibil public la: **%s**." % (DEMO_URL, REPO_URL)),

    ("h1", "2. Descrierea tehnologiilor utilizate"),
    ("p", "Proiectul folosește exclusiv tehnologii web standard, studiate în cadrul facultății, "
          "împreună cu instrumente moderne de versionare și publicare:"),
    ("bullets", [
        "**HTML5** — structura semantică a paginilor (header, main, section, article, footer, form).",
        "**CSS3** — designul vizual: variabile CSS (:root), Flexbox și CSS Grid pentru layout, "
        "animații și tranziții, media queries pentru un design 100% responsive (mobile-first).",
        "**JavaScript Vanilla (ES6+)** — toată interactivitatea: generarea navbar-ului și a footer-ului, "
        "filtrare și căutare, coșul de comandă, favoritele, validările formularelor și galeria cu lightbox.",
        "**localStorage** — salvarea datelor în browser (coș, favorite, rezervări), astfel încât acestea "
        "să persiste la reîncărcarea paginii, fără server și fără bază de date.",
        "**Git și GitHub** — versionarea codului și găzduirea publică a depozitului.",
        "**Vercel** — publicarea (deploy) site-ului static, cu auto-deploy la fiecare modificare din GitHub.",
    ]),
    ("p", "Nu au fost folosite biblioteci sau framework-uri externe (fără React, Vue, jQuery, Bootstrap "
          "sau Tailwind) și niciun CDN. Imaginile sunt fotografii reale cu licență liberă (Unsplash), "
          "descărcate local, astfel încât site-ul să nu depindă de niciun link extern."),

    ("h1", "3. Prezentarea structurii și funcționalităților site-ului"),
    ("h2", "3.1. Organizarea fișierelor"),
    ("p", "Codul este organizat pe fișiere separate, după responsabilități (HTML pentru structură, "
          "CSS pentru stil, JavaScript pentru logică):"),
    ("code",
     ".\n"
     "├── index.html            # Home\n"
     "├── meniu.html            # Meniu interactiv (filtrare, căutare, coș)\n"
     "├── evenimente.html       # Evenimente + favorite\n"
     "├── rezervare.html        # Formular rezervare cu validări\n"
     "├── galerie.html          # Galerie cu lightbox\n"
     "├── contact.html          # Formular contact cu validări\n"
     "├── css/\n"
     "│   └── style.css         # Design system + componente + responsive\n"
     "├── js/\n"
     "│   ├── main.js           # Navbar, footer, helpers localStorage, animații\n"
     "│   ├── data.js           # Datele: produse + evenimente\n"
     "│   ├── meniu.js          # Filtrare, căutare, coș, total\n"
     "│   ├── evenimente.js     # Render evenimente + favorite\n"
     "│   ├── rezervare.js      # Validări formular rezervare\n"
     "│   ├── contact.js        # Validări formular contact\n"
     "│   └── galerie.js        # Lightbox\n"
     "└── assets/images/        # Fotografii și grafică (logo, hartă)",
     "Structura folderelor proiectului"),

    ("h2", "3.2. Pagina Home (index.html)"),
    ("p", "Pagina de start cuprinde: o secțiune *hero* cu titlu, subtitlu și butoane de acțiune; "
          "secțiunea „Despre LUMINA”; trei carduri „De ce LUMINA”; o bandă de statistici cu numere "
          "animate (count-up); un preview cu trei produse populare (generate din datele aplicației) "
          "și un card cu următorul eveniment, calculat dinamic ca fiind cel mai apropiat eveniment viitor. "
          "Elementele apar treptat, cu un efect *fade-in* la derulare."),

    ("h2", "3.3. Pagina Meniu (meniu.html) — pagina centrală"),
    ("bullets", [
        "Produsele sunt încărcate din fișierul de date (13 produse în 4 categorii).",
        "**Filtrare pe categorii** (Cafea, Ceai, Deserturi, Brunch; „Toate” afișează tot).",
        "**Căutare live** după nume, fără diferență între majuscule și minuscule.",
        "Filtrarea și căutarea funcționează **combinat**.",
        "**Coș demo** cu cantități (+/−), ștergere produs și **total calculat automat** (preț × cantitate).",
        "**Persistență în localStorage**: coșul rămâne completat după reîncărcare; există butoane "
        "„Golește comanda” și „Trimite comanda (demo)”.",
        "Stări goale prietenoase (coș gol, căutare fără rezultate).",
    ]),

    ("h2", "3.4. Pagina Evenimente (evenimente.html)"),
    ("p", "Afișează evenimentele sortate cronologic. Fiecare eveniment are un buton de tip *toggle* "
          "„Adaugă la favorite / ♥ Favorit”, iar favoritele se salvează în localStorage și rămân marcate "
          "după reîncărcare. Un filtru „Toate / Doar favorite” și un contor completează funcționalitatea."),

    ("h2", "3.5. Pagina Rezervare (rezervare.html)"),
    ("p", "Conține un formular cu câmpurile: nume, email, telefon, dată, oră, număr de persoane și un "
          "mesaj opțional. **Validările sunt realizate în JavaScript** (nu doar prin atributul HTML "
          "„required”): nume de minim 2 caractere, email valid (expresie regulată), telefon cu minim 10 "
          "cifre, dată care nu poate fi în trecut și număr de persoane între 1 și 12. Mesajele de eroare "
          "apar clar sub fiecare câmp, iar la trimiterea validă se afișează un mesaj de succes."),

    ("h2", "3.6. Pagina Galerie (galerie.html)"),
    ("p", "Un grid responsive de imagini și un **lightbox custom**, scris de la zero în JavaScript: "
          "click pe imagine deschide un overlay pe tot ecranul, cu navigare înainte/înapoi (săgeți și "
          "tastele ← →), închidere cu „X”, click pe fundal sau tasta Esc. Cât timp lightbox-ul este "
          "deschis, derularea paginii din spate este blocată."),

    ("h2", "3.7. Pagina Contact (contact.html)"),
    ("p", "Un formular de contact (nume, email, subiect, mesaj) cu validări în JavaScript (mesaj de minim "
          "10 caractere) și o coloană cu informații utile: adresă (str. Eugen Doga 5, Chișinău — zonă "
          "pietonală), telefon, email, program și o hartă stilizată cu locația."),

    ("h2", "3.8. Componente comune și finisaje"),
    ("bullets", [
        "Navbar responsive, cu meniu „hamburger” animat pe mobil și link-ul paginii curente evidențiat.",
        "Footer identic pe toate paginile, cu anul curent generat dinamic și mențiunea autorului.",
        "Bară de progres la derulare, buton „înapoi sus” și animații *fade-in* în cascadă.",
        "Atenție la accesibilitate: contrast bun, texte alternative la imagini, focus vizibil, "
        "navigare cu tastatura pentru lightbox și meniul mobil.",
    ]),

    ("pagebreak",),
    ("h1", "4. Capturi de ecran"),
    ("p", "În continuare sunt prezentate capturile principalelor pagini ale aplicației, surprinse de pe "
          "varianta publicată online."),
    ("image", "home.png", "Figura 1 — Pagina Home (secțiunea hero)", 15.5),
    ("image", "meniu.png", "Figura 2 — Pagina Meniu, cu filtre, produse și coșul de comandă", 15.5),
    ("image", "evenimente.png", "Figura 3 — Pagina Evenimente, cu butoane de favorite", 15.5),
    ("image", "rezervare.png", "Figura 4 — Pagina Rezervare, cu formularul de rezervare", 15.5),
    ("image", "galerie.png", "Figura 5 — Pagina Galerie", 15.5),
    ("image", "contact.png", "Figura 6 — Pagina Contact, cu informații și hartă", 15.5),
    ("image", "mobil.png", "Figura 7 — Varianta pe mobil (design responsive)", 7.0),

    ("pagebreak",),
    ("h1", "5. Explicații pentru fiecare captură de ecran"),
    ("p", "**Figura 1 — Home.** Captura arată secțiunea *hero*, cu mesajul principal și butoanele „Vezi "
          "meniul” și „Rezervă o masă”, alături de o fotografie reală a unui interior de cafenea. "
          "Cuvintele-cheie din titlu sunt evidențiate în culoarea de accent, iar în partea de jos apare "
          "un indicator animat de derulare. Mai jos (în afara cadrului) se găsesc secțiunile despre, "
          "cardurile de prezentare și banda de statistici cu numere animate."),
    ("p", "**Figura 2 — Meniu.** Se observă butoanele de filtrare pe categorii, câmpul de căutare live și "
          "grila de produse cu fotografii reale. În dreapta este coșul de comandă, cu butoane +/− pentru "
          "cantitate și totalul calculat automat. Conținutul coșului se salvează în localStorage, deci "
          "rămâne completat și după reîncărcarea paginii."),
    ("p", "**Figura 3 — Evenimente.** Fiecare eveniment este prezentat ca un card cu imagine, dată, "
          "descriere și un buton de favorite. Apăsarea butonului marchează evenimentul (inima se umple), "
          "iar starea se păstrează în localStorage. Contorul din partea de sus arată câte favorite sunt."),
    ("p", "**Figura 4 — Rezervare.** Captura prezintă formularul de rezervare. La trimitere, fiecare câmp "
          "este verificat în JavaScript; câmpurile invalide sunt marcate vizual, iar mesajele de eroare "
          "apar sub ele. La completarea corectă se afișează un mesaj de confirmare."),
    ("p", "**Figura 5 — Galerie.** Grila de imagini este afișată responsive. La click pe o imagine se "
          "deschide lightbox-ul pe tot ecranul, cu imaginea mărită, săgeți de navigare, o legendă și un "
          "contor „x / y”. Lightbox-ul poate fi controlat și de la tastatură."),
    ("p", "**Figura 6 — Contact.** Captura arată formularul de contact alături de coloana cu informații "
          "(adresă în zona pietonală din centrul Chișinăului, telefon, email, program) și o hartă "
          "stilizată cu un marcaj de locație. Și aici validările sunt realizate în JavaScript."),
    ("p", "**Figura 7 — Mobil.** Ilustrează comportamentul responsive: navbar-ul devine meniu „hamburger”, "
          "iar conținutul se rearanjează pe o singură coloană, fără derulare orizontală, rămânând lizibil "
          "și pe ecrane mici."),

    ("pagebreak",),
    ("h1", "6. Concluzii"),
    ("p", "Proiectul LUMINA demonstrează că se poate construi un site modern, atractiv și complet "
          "funcțional folosind exclusiv tehnologiile web de bază, fără framework-uri și fără backend. "
          "Aplicația rulează prin simpla deschidere a fișierelor și a putut fi publicată gratuit pe o "
          "platformă de găzduire statică."),
    ("h2", "Ce am învățat (competențe dezvoltate)"),
    ("bullets", [
        "Structurarea semantică și accesibilă a paginilor în HTML5.",
        "Crearea unui *design system* consecvent în CSS3 (variabile, Grid, Flexbox, animații) și "
        "implementarea unui layout complet responsive, mobile-first.",
        "Programarea interactivității în JavaScript Vanilla: manipularea DOM-ului, evenimente, "
        "filtrare/căutare, lucrul cu array-uri și obiecte.",
        "Folosirea localStorage pentru persistarea datelor fără backend.",
        "Implementarea validărilor de formulare și a unei componente complexe (lightbox) de la zero.",
        "Organizarea codului pe fișiere, eliminarea duplicării prin funcții comune și comentarea clară a logicii.",
        "Versionarea cu Git/GitHub și publicarea unui site prin Vercel, cu auto-deploy.",
    ]),
    ("p", "Rezultatul final este o aplicație web coerentă, ușor de explicat și de extins, care reflectă "
          "cunoștințele acumulate în primul an de studiu."),

    ("pagebreak",),
    ("h1", "7. Anexă — fragmente relevante din codul sursă"),
    ("p", "Mai jos sunt prezentate fragmentele cele mai importante din cod, comentate, organizate pe teme. "
          "Codul complet se găsește în depozitul proiectului."),

    ("h2", "7.1. Helpers pentru localStorage (js/main.js)"),
    ("p", "Centralizează citirea/scrierea în localStorage, cu transformare automată în/din JSON. Sunt "
          "refolosite de toate paginile, ca să nu existe cod duplicat."),
    ("code",
     "// Citește o valoare din localStorage și o transformă din JSON.\n"
     "function citesteStocare(cheie, implicit) {\n"
     "  try {\n"
     "    const brut = localStorage.getItem(cheie);\n"
     "    return brut ? JSON.parse(brut) : implicit;\n"
     "  } catch (e) {\n"
     "    return implicit;\n"
     "  }\n"
     "}\n"
     "// Scrie o valoare în localStorage (o serializează în JSON).\n"
     "function scrieStocare(cheie, valoare) {\n"
     "  localStorage.setItem(cheie, JSON.stringify(valoare));\n"
     "}", None),

    ("h2", "7.2. Filtrare + căutare combinate (js/meniu.js)"),
    ("p", "Filtrează produsele simultan după categoria selectată și după textul căutat."),
    ("code",
     "function produseFiltrate() {\n"
     "  return PRODUSE.filter(function (p) {\n"
     "    // Categoria: „Toate” trece tot\n"
     "    const okCategorie = categorieActiva === 'Toate' || p.categorie === categorieActiva;\n"
     "    // Căutarea: numele conține textul (case-insensitive)\n"
     "    const okCautare = p.nume.toLowerCase().indexOf(textCautare.toLowerCase()) !== -1;\n"
     "    return okCategorie && okCautare;\n"
     "  });\n"
     "}", None),

    ("h2", "7.3. Calculul automat al totalului din coș (js/meniu.js)"),
    ("p", "Parcurge produsele din coș, înmulțește prețul cu cantitatea și adună totul."),
    ("code",
     "let total = 0;\n"
     "idsuri.forEach(function (idStr) {\n"
     "  const produs = gasesteProdus(parseInt(idStr, 10));\n"
     "  const cant = cos[idStr];\n"
     "  total += produs.pret * cant; // preț × cantitate\n"
     "});\n"
     "elCosTotal.textContent = total + ' lei';", None),

    ("h2", "7.4. Favorite salvate în localStorage (js/evenimente.js)"),
    ("p", "Adaugă sau scoate un eveniment din favorite (toggle) și salvează imediat starea."),
    ("code",
     "function comutaFavorit(id) {\n"
     "  if (esteFavorit(id)) {\n"
     "    favorite = favorite.filter(function (x) { return x !== id; });\n"
     "  } else {\n"
     "    favorite.push(id);\n"
     "  }\n"
     "  scrieStocare(CHEIE_FAV, favorite); // salvăm imediat\n"
     "  randeaza();\n"
     "}", None),

    ("h2", "7.5. Validarea formularului de rezervare (js/rezervare.js)"),
    ("p", "Validări în JavaScript, cu mesaje de eroare clare. Exemple: email cu expresie regulată și "
          "dată care nu poate fi în trecut."),
    ("code",
     "// Email – format valid\n"
     "const regexEmail = /^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$/;\n"
     "if (!regexEmail.test(email)) {\n"
     "  seteazaEroare('email', 'err-email', 'Introdu o adresă de email validă.');\n"
     "  valid = false;\n"
     "}\n"
     "// Data – să nu fie în trecut\n"
     "const aleasa = new Date(data);\n"
     "const azi = new Date(); azi.setHours(0, 0, 0, 0);\n"
     "if (aleasa < azi) {\n"
     "  seteazaEroare('data', 'err-data', 'Data nu poate fi în trecut.');\n"
     "  valid = false;\n"
     "}", None),

    ("h2", "7.6. Lightbox cu navigare ciclică (js/galerie.js)"),
    ("p", "Operatorul modulo (%) face navigarea să „treacă” de la ultima imagine la prima și invers."),
    ("code",
     "function imagineUrmatoare() {\n"
     "  indexCurent = (indexCurent + 1) % IMAGINI.length;\n"
     "  actualizeazaImagine();\n"
     "}\n"
     "// Navigare cu tastatura, doar când lightbox-ul e deschis\n"
     "document.addEventListener('keydown', function (e) {\n"
     "  if (!lightbox.classList.contains('deschis')) return;\n"
     "  if (e.key === 'Escape') inchideLightbox();\n"
     "  else if (e.key === 'ArrowRight') imagineUrmatoare();\n"
     "  else if (e.key === 'ArrowLeft') imagineAnterioara();\n"
     "});", None),

    ("h2", "7.7. Animație fade-in la derulare (js/main.js)"),
    ("p", "Folosește IntersectionObserver pentru a anima elementele când intră în ecran."),
    ("code",
     "const observator = new IntersectionObserver(function (intrari) {\n"
     "  intrari.forEach(function (intrare) {\n"
     "    if (intrare.isIntersecting) {\n"
     "      intrare.target.classList.add('vizibil');\n"
     "      observator.unobserve(intrare.target); // animăm o singură dată\n"
     "    }\n"
     "  });\n"
     "}, { threshold: 0.12 });", None),

    ("h2", "7.8. Variabile de design în CSS (css/style.css)"),
    ("p", "Paleta și spacing-ul sunt definite o singură dată în :root și refolosite peste tot."),
    ("code",
     ":root {\n"
     "  --crem: #F7F0E6;\n"
     "  --caramel: #C8893F;\n"
     "  --maro: #3B2A20;\n"
     "  --masliniu: #6B7A4F;\n"
     "  --sp-3: 24px;\n"
     "  --radius: 16px;\n"
     "  --tranzitie: 250ms ease;\n"
     "}", None),
]


# ======================================================================
# 2. RENDERER MARKDOWN
# ======================================================================
def render_md():
    L = []
    c = COVER
    L.append("# %s\n" % c["titlu"])
    L.append("### %s\n" % c["subtitlu"])
    L.append("")
    L.append("**%s**  " % c["supratitlu"])
    L.append("")
    L.append("**Autor:** %s  " % c["autor"])
    L.append("**%s**  " % c["grupa"])
    L.append("**Specializare:** %s, Facultatea de Matematică și Informatică  " % c["an"])
    L.append("**%s**  " % c["loc"])
    L.append("")
    L.append("**Demo online:** %s  " % c["demo"])
    L.append("**Cod sursă:** %s  " % REPO_URL)
    L.append("\n---\n")
    L.append("> *Cuprinsul este generat automat în varianta Word/PDF a raportului.*\n")
    L.append("\n---\n")

    for bloc in CONTENT:
        t = bloc[0]
        if t == "h1":
            L.append("\n## %s\n" % bloc[1])
        elif t == "h2":
            L.append("\n### %s\n" % bloc[1])
        elif t == "h3":
            L.append("\n#### %s\n" % bloc[1])
        elif t == "p":
            L.append(bloc[1] + "\n")
        elif t == "bullets":
            for it in bloc[1]:
                L.append("- " + it)
            L.append("")
        elif t == "code":
            cap = bloc[2]
            if cap:
                L.append("*%s*\n" % cap)
            L.append("```")
            L.append(bloc[1])
            L.append("```\n")
        elif t == "image":
            fn, cap = bloc[1], bloc[2]
            L.append("![%s](../assets/images/screenshots/%s)" % (cap, fn))
            L.append("*%s*\n" % cap)
        elif t == "pagebreak":
            L.append("\n---\n")
    with open(MD_OUT, "w", encoding="utf-8") as f:
        f.write("\n".join(L))
    print("OK  Markdown ->", MD_OUT)


# ======================================================================
# 3. RENDERER WORD (.docx)
# ======================================================================
def set_cell_run(run, text, size=11, bold=False, color=None, font="Segoe UI"):
    run.text = text
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font
    if color:
        run.font.color.rgb = color


def add_runs_bold(paragraph, text, size=11, color=None, font="Segoe UI"):
    """Adaugă text într-un paragraf, interpretând **bold** între asteriscuri."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        r = paragraph.add_run(part)
        r.font.size = Pt(size)
        r.font.name = font
        r.font.bold = (i % 2 == 1)
        if color:
            r.font.color.rgb = color


def shade_paragraph(paragraph, fill="F1E9DC"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    for typ, txt in [("begin", None), (None, "PAGE"), ("end", None)]:
        if typ:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), typ)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            run._r.append(instr)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Actualizați câmpul (clic dreapta → Update Field / F9)."
    fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "end")
    for el in (fld1, instr, fld2, t, fld3):
        run._r.append(el)


def render_docx():
    doc = Document()

    # Stil implicit
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)
    normal.font.color.rgb = MARO_MEDIU
    doc.styles["Normal"].paragraph_format.space_after = Pt(8)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.25

    # Margini
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    # Stiluri de titlu
    for name, size, color in [("Heading 1", 17, MARO), ("Heading 2", 13, CARAMEL_INCHIS)]:
        st = doc.styles[name]
        st.font.name = "Georgia"; st.font.size = Pt(size)
        st.font.color.rgb = color; st.font.bold = True
        st.paragraph_format.space_before = Pt(14); st.paragraph_format.space_after = Pt(6)

    # --- PAGINA DE TITLU ---
    c = COVER
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_run(p.add_run(), c["supratitlu"], size=12, bold=True, color=CARAMEL_INCHIS)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    set_cell_run(p.add_run(), c["titlu"], size=34, bold=True, color=MARO, font="Georgia")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_run(p.add_run(), c["subtitlu"], size=15, bold=False, color=MARO_MEDIU, font="Georgia")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    set_cell_run(p.add_run(), "Autor:  ", size=13, bold=False, color=MARO_MEDIU)
    set_cell_run(p.add_run(), c["autor"], size=13, bold=True, color=MARO)

    for line in [c["grupa"], "%s · Facultatea de Matematică și Informatică" % c["an"], c["loc"]]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_run(p.add_run(), line, size=12, bold=False, color=MARO_MEDIU)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    set_cell_run(p.add_run(), "Demo online:  ", size=10, color=MARO_MEDIU)
    set_cell_run(p.add_run(), c["demo"], size=10, bold=True, color=CARAMEL_INCHIS)

    # --- CUPRINS ---
    doc.add_page_break()
    p = doc.add_paragraph()
    set_cell_run(p.add_run(), "Cuprins", size=17, bold=True, color=MARO, font="Georgia")
    add_toc(doc)
    doc.add_page_break()

    # --- CONȚINUT ---
    for bloc in CONTENT:
        t = bloc[0]
        if t == "h1":
            doc.add_paragraph(bloc[1], style="Heading 1")
        elif t == "h2":
            doc.add_paragraph(bloc[1], style="Heading 2")
        elif t == "h3":
            doc.add_paragraph(bloc[1], style="Heading 2")
        elif t == "p":
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs_bold(p, bloc[1])
        elif t == "bullets":
            for it in bloc[1]:
                p = doc.add_paragraph(style="List Bullet")
                add_runs_bold(p, it)
        elif t == "code":
            cap = bloc[2]
            if cap:
                pc = doc.add_paragraph()
                r = pc.add_run(cap); r.font.italic = True; r.font.size = Pt(9.5)
                r.font.color.rgb = MARO_MEDIU
            for line in bloc[1].split("\n"):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                shade_paragraph(p)
                r = p.add_run(line if line else " ")
                r.font.name = "Consolas"; r.font.size = Pt(9)
                r.font.color.rgb = MARO
            doc.add_paragraph()
        elif t == "image":
            fn, cap, width = bloc[1], bloc[2], bloc[3]
            path = os.path.join(SHOTS, fn)
            if os.path.exists(path):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Cm(width))
            else:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                shade_paragraph(p)
                set_cell_run(p.add_run(), "[Captură de inserat: %s]" % fn, size=10, color=MARO_MEDIU)
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.add_run(cap); r.font.italic = True; r.font.size = Pt(9.5)
            r.font.color.rgb = MARO_MEDIU
        elif t == "pagebreak":
            doc.add_page_break()

    # Numerotarea paginilor în subsol
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

    doc.save(DOCX_OUT)
    print("OK  Word     ->", DOCX_OUT)


# ======================================================================
# 4. CONVERSIE PDF (prin Microsoft Word / win32com) + actualizare cuprins
# ======================================================================
def render_pdf():
    try:
        import win32com.client
    except ImportError:
        print("!!  win32com indisponibil – sar peste PDF.")
        return
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(DOCX_OUT)
        # Actualizează cuprinsul (TOC) și toate câmpurile
        try:
            doc.TablesOfContents(1).Update()
        except Exception:
            pass
        for story in doc.StoryRanges:
            story.Fields.Update()
        doc.Save()
        doc.SaveAs(PDF_OUT, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close(False)
        print("OK  PDF      ->", PDF_OUT)
    finally:
        word.Quit()


if __name__ == "__main__":
    render_md()
    render_docx()
    render_pdf()
    print("\nGata. Toate fisierele au fost generate in folderul docs/.")
